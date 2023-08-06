import io
from pathlib import Path
from typing import Any, Union

import docx
import docx.document
from docxtpl import DocxTemplate


def render_docx_from_template(docx_path: Union[str, Path], **context: Any) -> bytes:
    """
    Рендерит docx из шаблона {docx_path}
    :param docx_path: Путь к docx-шаблону - docx-файл, использующий jinja-like синтаксис согласно `docxtpl <https://docxtpl.readthedocs.io/en/latest/>`_
    :param context: kwargs, которые передаются в шаблон (аналогично тому, как это делается в jinja)
    :return: docx-байты
    """
    doc = DocxTemplate(docx_path)
    doc.render(context)
    doc.save(stream := io.BytesIO())
    return stream.getvalue()


def docxs_are_equal(docx_1: bytes, docx_2: bytes) -> bool:
    def get_paragraph_data(doc: docx.document.Document) -> Any:
        return tuple(p.text for p in doc.paragraphs)

    def get_table_data(doc: docx.document.Document) -> Any:
        return tuple(
            tuple(cell.text for cell in row.cells)
            for table in doc.tables for row in table.rows
        )

    docx_1_data = docx.Document(io.BytesIO(docx_1))
    docx_2_data = docx.Document(io.BytesIO(docx_2))

    return (
        get_paragraph_data(docx_1_data) == get_paragraph_data(docx_2_data) and
        get_table_data(docx_1_data) == get_table_data(docx_2_data)
    )
